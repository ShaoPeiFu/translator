#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF处理模块
支持PDF文本提取、翻译和导出功能
"""

import os
import re
from typing import Dict, List, Optional, Tuple
from pathlib import Path
import PyPDF2
import pdfplumber
from PIL import Image
import io
import base64


class PDFProcessor:
    """PDF文档处理器"""

    def __init__(self):
        """初始化PDF处理器"""
        self.supported_formats = [".pdf", ".docx", ".txt"]

    def extract_text_from_pdf(self, pdf_path: str) -> Dict[str, any]:
        """
        从PDF文件提取文本

        Args:
            pdf_path: PDF文件路径

        Returns:
            包含提取结果的字典
        """
        try:
            text_content = []
            page_info = []

            # 使用pdfplumber提取文本（更好的文本保持）
            with pdfplumber.open(pdf_path) as pdf:
                total_pages = len(pdf.pages)

                for page_num, page in enumerate(pdf.pages, 1):
                    # 提取文本
                    text = page.extract_text()
                    if text:
                        # 清理文本
                        cleaned_text = self._clean_text(text)
                        text_content.append(cleaned_text)

                        page_info.append(
                            {
                                "page": page_num,
                                "text": cleaned_text,
                                "char_count": len(cleaned_text),
                                "word_count": len(cleaned_text.split()),
                            }
                        )
                    else:
                        # 如果无法提取文本，可能是扫描版PDF
                        text_content.append(
                            f"[第{page_num}页 - 无法提取文本，可能是扫描版PDF]"
                        )
                        page_info.append(
                            {
                                "page": page_num,
                                "text": "",
                                "char_count": 0,
                                "word_count": 0,
                                "note": "可能是扫描版PDF",
                            }
                        )

            return {
                "status": "success",
                "total_pages": total_pages,
                "total_text": "\n\n".join(text_content),
                "page_info": page_info,
                "file_path": pdf_path,
                "file_name": os.path.basename(pdf_path),
            }

        except Exception as e:
            return {"status": "error", "error": str(e), "file_path": pdf_path}

    def _clean_text(self, text: str) -> str:
        """
        清理提取的文本

        Args:
            text: 原始文本

        Returns:
            清理后的文本
        """
        if not text:
            return ""

        # 移除多余的空白字符
        text = re.sub(r"\s+", " ", text)

        # 移除页眉页脚（通常包含页码等）
        text = re.sub(r"^\s*\d+\s*$", "", text, flags=re.MULTILINE)

        # 移除孤立的字符
        text = re.sub(r"\s+[a-zA-Z]\s+", " ", text)

        # 清理段落分隔
        text = re.sub(r"\n\s*\n", "\n\n", text)

        return text.strip()

    def split_text_for_translation(
        self, text: str, max_chunk_size: int = 1000
    ) -> List[str]:
        """
        将长文本分割成适合翻译的块

        Args:
            text: 要分割的文本
            max_chunk_size: 每个块的最大字符数

        Returns:
            文本块列表
        """
        if len(text) <= max_chunk_size:
            return [text]

        chunks = []
        sentences = re.split(r"[.!?。！？]\s*", text)

        current_chunk = ""
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue

            # 如果当前块加上新句子超过限制，保存当前块并开始新块
            if len(current_chunk) + len(sentence) > max_chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = sentence
            else:
                if current_chunk:
                    current_chunk += ". " + sentence
                else:
                    current_chunk = sentence

        # 添加最后一个块
        if current_chunk:
            chunks.append(current_chunk.strip())

        return chunks

    def merge_translated_chunks(
        self, original_chunks: List[str], translated_chunks: List[str]
    ) -> str:
        """
        合并翻译后的文本块

        Args:
            original_chunks: 原始文本块列表
            translated_chunks: 翻译后的文本块列表

        Returns:
            合并后的完整翻译文本
        """
        if len(original_chunks) != len(translated_chunks):
            return "翻译块数量不匹配，请检查翻译结果"

        merged_text = ""
        for i, (orig, trans) in enumerate(zip(original_chunks, translated_chunks)):
            if i > 0:
                merged_text += "\n\n"
            merged_text += f"原文 (块 {i+1}):\n{orig}\n\n翻译:\n{trans}"

        return merged_text

    def create_comparison_pdf(
        self,
        original_text: str,
        translated_text: str,
        output_path: str,
        title: str = "翻译对比",
    ) -> Dict[str, str]:
        """
        创建原文和翻译的对比PDF

        Args:
            original_text: 原文
            translated_text: 翻译文本
            output_path: 输出PDF路径
            title: PDF标题

        Returns:
            创建结果
        """
        try:
            # 检查输出目录权限
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                try:
                    os.makedirs(output_dir, exist_ok=True)
                except PermissionError:
                    return {
                        "status": "error",
                        "error": f"无法创建输出目录: {output_dir}",
                    }
                except Exception as e:
                    return {"status": "error", "error": f"创建输出目录失败: {str(e)}"}

            # 检查文件写入权限
            try:
                with open(output_path, "w") as f:
                    pass
                os.remove(output_path)
            except PermissionError:
                return {"status": "error", "error": f"没有写入权限: {output_path}"}
            except Exception as e:
                return {"status": "error", "error": f"检查文件权限失败: {str(e)}"}

            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import (
                SimpleDocTemplate,
                Paragraph,
                Spacer,
                PageBreak,
            )
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont

            # 改进的字体注册逻辑
            chinese_font = "Helvetica"  # 默认字体

            try:
                # 对于macOS系统 - 尝试多个可能的字体路径
                mac_fonts = [
                    "/System/Library/Fonts/PingFang.ttc",
                    "/System/Library/Fonts/STHeiti Light.ttc",
                    "/System/Library/Fonts/STHeiti Medium.ttc",
                    "/Library/Fonts/Arial Unicode MS.ttf",
                    "/System/Library/Fonts/Helvetica.ttc",
                ]

                for font_path in mac_fonts:
                    if os.path.exists(font_path):
                        try:
                            if font_path.endswith(".ttc"):
                                # 对于.ttc字体文件，尝试注册
                                pdfmetrics.registerFont(
                                    TTFont("ChineseFont", font_path)
                                )
                                chinese_font = "ChineseFont"
                                break
                            elif font_path.endswith(".ttf"):
                                pdfmetrics.registerFont(
                                    TTFont("ChineseFont", font_path)
                                )
                                chinese_font = "ChineseFont"
                                break
                        except Exception:
                            continue

            except Exception:
                # 如果字体注册失败，使用默认字体
                chinese_font = "Helvetica"

            # 创建PDF文档
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            story = []

            # 定义样式
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                "CustomTitle",
                parent=styles["Heading1"],
                fontSize=18,
                spaceAfter=30,
                alignment=1,  # 居中
                fontName=chinese_font,
            )

            subtitle_style = ParagraphStyle(
                "CustomSubtitle",
                parent=styles["Heading2"],
                fontSize=14,
                spaceAfter=20,
                fontName=chinese_font,
            )

            content_style = ParagraphStyle(
                "CustomContent",
                parent=styles["Normal"],
                fontSize=10,
                spaceAfter=12,
                fontName=chinese_font,
                leading=14,
            )

            # 添加标题
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 20))

            # 添加原文 - 显示完整内容，不限制段落数量
            story.append(Paragraph("原文 (Original Text)", subtitle_style))
            story.append(Spacer(1, 10))

            # 处理原文，显示所有段落
            original_paragraphs = original_text.split("\n\n")
            for para in original_paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), content_style))
                    story.append(Spacer(1, 8))

            story.append(PageBreak())

            # 添加翻译 - 显示完整内容，不限制段落数量
            story.append(Paragraph("翻译 (Translation)", subtitle_style))
            story.append(Spacer(1, 10))

            # 处理翻译文本，显示所有段落
            translated_paragraphs = translated_text.split("\n\n")
            for para in translated_paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), content_style))
                    story.append(Spacer(1, 8))

            # 生成PDF
            doc.build(story)

            # 验证文件是否成功创建
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                return {
                    "status": "success",
                    "output_path": output_path,
                    "format": "pdf",
                }
            else:
                return {"status": "error", "error": "PDF文件生成失败或文件为空"}

        except ImportError as e:
            return {"status": "error", "error": f"缺少必要的依赖库: {str(e)}"}
        except Exception as e:
            return {"status": "error", "error": f"PDF生成过程中出现错误: {str(e)}"}

    def create_enhanced_comparison_pdf(
        self,
        original_text: str,
        translated_text: str,
        output_path: str,
        title: str = "翻译对比",
        preserve_formatting: bool = True,
    ) -> Dict[str, str]:
        """
        创建增强的PDF对比，更好地保留格式和显示完整内容

        Args:
            original_text: 原文
            translated_text: 翻译文本
            output_path: 输出PDF路径
            title: PDF标题
            preserve_formatting: 是否保留格式

        Returns:
            创建结果
        """
        try:
            # 检查输出目录权限
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_path):
                try:
                    os.makedirs(output_dir, exist_ok=True)
                except PermissionError:
                    return {
                        "status": "error",
                        "error": f"无法创建输出目录: {output_dir}",
                    }
                except Exception as e:
                    return {"status": "error", "error": f"创建输出目录失败: {str(e)}"}

            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import (
                SimpleDocTemplate,
                Paragraph,
                Spacer,
                PageBreak,
                Table,
                TableStyle,
            )
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont

            # 改进的字体注册逻辑
            chinese_font = "Helvetica"

            try:
                # 对于macOS系统 - 尝试多个可能的字体路径
                mac_fonts = [
                    "/System/Library/Fonts/PingFang.ttc",
                    "/System/Library/Fonts/STHeiti Light.ttc",
                    "/System/Library/Fonts/STHeiti Medium.ttc",
                    "/Library/Fonts/Arial Unicode MS.ttf",
                    "/System/Library/Fonts/Helvetica.ttc",
                ]

                for font_path in mac_fonts:
                    if os.path.exists(font_path):
                        try:
                            if font_path.endswith(".ttc"):
                                pdfmetrics.registerFont(
                                    TTFont("ChineseFont", font_path)
                                )
                                chinese_font = "ChineseFont"
                                break
                            elif font_path.endswith(".ttf"):
                                pdfmetrics.registerFont(
                                    TTFont("ChineseFont", font_path)
                                )
                                chinese_font = "ChineseFont"
                                break
                        except Exception:
                            continue

            except Exception:
                chinese_font = "Helvetica"

            # 创建PDF文档
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            story = []

            # 定义样式
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                "CustomTitle",
                parent=styles["Heading1"],
                fontSize=18,
                spaceAfter=30,
                alignment=1,
                fontName=chinese_font,
            )

            subtitle_style = ParagraphStyle(
                "CustomSubtitle",
                parent=styles["Heading2"],
                fontSize=14,
                spaceAfter=20,
                fontName=chinese_font,
            )

            content_style = ParagraphStyle(
                "CustomContent",
                parent=styles["Normal"],
                fontSize=10,
                spaceAfter=12,
                fontName=chinese_font,
                leading=14,
            )

            # 添加标题
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 20))

            if preserve_formatting:
                # 使用表格格式进行对比
                story.append(Paragraph("原文与翻译对比", subtitle_style))
                story.append(Spacer(1, 15))

                # 分割文本为段落
                original_paragraphs = [
                    p.strip() for p in original_text.split("\n\n") if p.strip()
                ]
                translated_paragraphs = [
                    p.strip() for p in translated_text.split("\n\n") if p.strip()
                ]

                # 确保两个列表长度一致
                max_paragraphs = max(
                    len(original_paragraphs), len(translated_paragraphs)
                )
                original_paragraphs.extend(
                    [""] * (max_paragraphs - len(original_paragraphs))
                )
                translated_paragraphs.extend(
                    [""] * (max_paragraphs - len(translated_paragraphs))
                )

                # 创建对比表格
                table_data = []
                table_data.append(
                    [
                        Paragraph("原文 (Original)", subtitle_style),
                        Paragraph("翻译 (Translation)", subtitle_style),
                    ]
                )

                # 添加所有段落到表格
                for i, (orig, trans) in enumerate(
                    zip(original_paragraphs, translated_paragraphs)
                ):
                    if orig or trans:  # 只添加非空段落
                        table_data.append(
                            [
                                Paragraph(orig, content_style),
                                Paragraph(trans, content_style),
                            ]
                        )

                # 创建表格
                table = Table(
                    table_data, colWidths=[doc.width / 2 - 20, doc.width / 2 - 20]
                )
                table.setStyle(
                    TableStyle(
                        [
                            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("FONTNAME", (0, 0), (-1, 0), chinese_font),
                            ("FONTSIZE", (0, 0), (-1, 0), 12),
                            ("FONTNAME", (0, 1), (-1, -1), chinese_font),
                            ("FONTSIZE", (0, 1), (-1, -1), 9),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                            ("TOPPADDING", (0, 0), (-1, -1), 6),
                            ("LEFTPADDING", (0, 0), (-1, -1), 10),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                            ("GRID", (0, 0), (-1, -1), 1, colors.grey),
                            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                        ]
                    )
                )

                story.append(table)
            else:
                # 使用分页格式
                # 添加原文
                story.append(Paragraph("原文 (Original Text)", subtitle_style))
                story.append(Spacer(1, 10))

                # 处理原文，显示所有段落
                original_paragraphs = original_text.split("\n\n")
                for para in original_paragraphs:
                    if para.strip():
                        story.append(Paragraph(para.strip(), content_style))
                        story.append(Spacer(1, 8))

                story.append(PageBreak())

                # 添加翻译
                story.append(Paragraph("翻译 (Translation)", subtitle_style))
                story.append(Spacer(1, 10))

                # 处理翻译文本，显示所有段落
                translated_paragraphs = translated_text.split("\n\n")
                for para in translated_paragraphs:
                    if para.strip():
                        story.append(Paragraph(para.strip(), content_style))
                        story.append(Spacer(1, 8))

            # 生成PDF
            doc.build(story)

            # 验证文件是否成功创建
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                return {
                    "status": "success",
                    "output_path": output_path,
                    "format": "pdf",
                }
            else:
                return {"status": "error", "error": "PDF文件生成失败或文件为空"}

        except ImportError as e:
            return {"status": "error", "error": f"缺少必要的依赖库: {str(e)}"}
        except Exception as e:
            return {"status": "error", "error": f"PDF生成过程中出现错误: {str(e)}"}

    def create_mineru_optimized_pdf(
        self,
        original_text: str,
        translated_text: str,
        output_path: str,
        title: str = "翻译对比 - MinerU优化排版",
    ) -> Dict[str, str]:
        """
        使用MinerU排版优化创建PDF对比

        Args:
            original_text: 原文
            translated_text: 翻译文本
            output_path: 输出PDF路径
            title: PDF标题

        Returns:
            创建结果
        """
        try:
            # 检查输出目录权限
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                try:
                    os.makedirs(output_dir, exist_ok=True)
                except PermissionError:
                    return {
                        "status": "error",
                        "error": f"无法创建输出目录: {output_dir}",
                    }
                except Exception as e:
                    return {"status": "error", "error": f"创建输出目录失败: {str(e)}"}

            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import (
                SimpleDocTemplate,
                Paragraph,
                Spacer,
                PageBreak,
                Table,
                TableStyle,
                Frame,
                PageTemplate,
                BaseDocTemplate,
            )
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch, cm
            from reportlab.lib import colors
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

            # 改进的字体注册逻辑
            chinese_font = "Helvetica"

            try:
                # 对于macOS系统 - 尝试多个可能的字体路径
                mac_fonts = [
                    "/System/Library/Fonts/PingFang.ttc",
                    "/System/Library/Fonts/STHeiti Light.ttc",
                    "/System/Library/Fonts/STHeiti Medium.ttc",
                    "/Library/Fonts/Arial Unicode MS.ttf",
                    "/System/Library/Fonts/Helvetica.ttc",
                ]

                for font_path in mac_fonts:
                    if os.path.exists(font_path):
                        try:
                            if font_path.endswith(".ttc"):
                                pdfmetrics.registerFont(
                                    TTFont("ChineseFont", font_path)
                                )
                                chinese_font = "ChineseFont"
                                break
                            elif font_path.endswith(".ttf"):
                                pdfmetrics.registerFont(
                                    TTFont("ChineseFont", font_path)
                                )
                                chinese_font = "ChineseFont"
                                break
                        except Exception:
                            continue

            except Exception:
                chinese_font = "Helvetica"

            # 创建PDF文档 - 使用SimpleDocTemplate以获得更多控制
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            story = []

            # 定义优化的样式
            styles = getSampleStyleSheet()

            # 主标题样式 - 更加突出
            title_style = ParagraphStyle(
                "EnhancedTitle",
                parent=styles["Heading1"],
                fontSize=22,
                spaceAfter=25,
                alignment=TA_CENTER,
                fontName=chinese_font,
                textColor=colors.darkblue,
                spaceBefore=15,
                borderWidth=1,
                borderColor=colors.darkblue,
                borderPadding=8,
                backColor=colors.lightblue,
            )

            # 副标题样式 - 更清晰
            subtitle_style = ParagraphStyle(
                "EnhancedSubtitle",
                parent=styles["Heading2"],
                fontSize=16,
                spaceAfter=20,
                spaceBefore=12,
                fontName=chinese_font,
                textColor=colors.darkgreen,
                borderWidth=0.5,
                borderColor=colors.darkgreen,
                borderPadding=5,
                backColor=colors.lightgreen,
            )

            # 内容样式 - 更易读
            content_style = ParagraphStyle(
                "EnhancedContent",
                parent=styles["Normal"],
                fontSize=11,
                spaceAfter=10,
                spaceBefore=5,
                fontName=chinese_font,
                leading=16,
                alignment=TA_LEFT,  # 使用左对齐避免justification问题
                textColor=colors.black,
                leftIndent=15,
                rightIndent=15,
            )

            # 原文样式 - 特殊标识
            original_style = ParagraphStyle(
                "OriginalStyle",
                parent=content_style,
                textColor=colors.darkred,
                backColor=colors.lightgrey,
                borderWidth=0.5,
                borderColor=colors.darkred,
                borderPadding=3,
                leftIndent=20,
            )

            # 翻译样式 - 特殊标识
            translated_style = ParagraphStyle(
                "TranslatedStyle",
                parent=content_style,
                textColor=colors.darkblue,
                backColor=colors.lightyellow,
                borderWidth=0.5,
                borderColor=colors.darkblue,
                borderPadding=3,
                leftIndent=20,
            )

            # 添加主标题
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 20))

            # 使用MinerU优化的布局
            # 分割文本为段落
            original_paragraphs = [
                p.strip() for p in original_text.split("\n\n") if p.strip()
            ]
            translated_paragraphs = [
                p.strip() for p in translated_text.split("\n\n") if p.strip()
            ]

            # 确保两个列表长度一致
            max_paragraphs = max(len(original_paragraphs), len(translated_paragraphs))
            original_paragraphs.extend(
                [""] * (max_paragraphs - len(original_paragraphs))
            )
            translated_paragraphs.extend(
                [""] * (max_paragraphs - len(translated_paragraphs))
            )

            # 创建优化的对比表格
            table_data = []

            # 表头 - 更加突出
            table_data.append(
                [
                    Paragraph(
                        "📖 原文 (Original)",
                        ParagraphStyle(
                            "TableHeader",
                            parent=subtitle_style,
                            alignment=TA_CENTER,
                            fontSize=14,
                            backColor=colors.darkred,
                            textColor=colors.white,
                            borderPadding=8,
                        ),
                    ),
                    Paragraph(
                        "🌐 翻译 (Translation)",
                        ParagraphStyle(
                            "TableHeader",
                            parent=subtitle_style,
                            alignment=TA_CENTER,
                            fontSize=14,
                            backColor=colors.darkblue,
                            textColor=colors.white,
                            borderPadding=8,
                        ),
                    ),
                ]
            )

            # 添加所有段落到表格，使用MinerU优化的样式
            for i, (orig, trans) in enumerate(
                zip(original_paragraphs, translated_paragraphs)
            ):
                if orig or trans:  # 只添加非空段落
                    # 为段落添加序号标识
                    orig_para = f"【{i+1}】{orig}" if orig else ""
                    trans_para = f"【{i+1}】{trans}" if trans else ""

                    table_data.append(
                        [
                            (
                                Paragraph(orig_para, original_style)
                                if orig
                                else Paragraph("", content_style)
                            ),
                            (
                                Paragraph(trans_para, translated_style)
                                if trans
                                else Paragraph("", content_style)
                            ),
                        ]
                    )

            # 创建表格 - 优化列宽和样式
            table = Table(
                table_data,
                colWidths=[doc.width / 2 - 25, doc.width / 2 - 25],
                repeatRows=1,  # 重复表头
            )

            # 优化的表格样式
            table.setStyle(
                TableStyle(
                    [
                        # 对齐和布局
                        ("ALIGN", (0, 0), (-1, -1), TA_LEFT),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        # 字体设置
                        ("FONTNAME", (0, 0), (-1, 0), chinese_font),
                        ("FONTSIZE", (0, 0), (-1, 0), 14),
                        ("FONTNAME", (0, 1), (-1, -1), chinese_font),
                        ("FONTSIZE", (0, 1), (-1, -1), 10),
                        # 间距设置
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                        ("TOPPADDING", (0, 0), (-1, -1), 8),
                        ("LEFTPADDING", (0, 0), (-1, -1), 12),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 12),
                        # 边框和背景
                        ("GRID", (0, 0), (-1, -1), 1, colors.grey),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.darkgrey),
                        (
                            "ROWBACKGROUNDS",
                            (0, 1),
                            (-1, -1),
                            [colors.white, colors.lightgrey],
                        ),
                        # 特殊行样式
                        ("LINEBELOW", (0, 0), (-1, 0), 2, colors.black),
                        ("LINEABOVE", (0, 0), (-1, 0), 2, colors.black),
                    ]
                )
            )

            story.append(table)

            # 添加统计信息
            story.append(Spacer(1, 20))
            info_style = ParagraphStyle(
                "InfoStyle",
                parent=content_style,
                fontSize=10,
                textColor=colors.grey,
                alignment=TA_CENTER,
            )
            story.append(
                Paragraph(
                    f"📊 对比统计：原文 {len([p for p in original_paragraphs if p])} 段，"
                    f"翻译 {len([p for p in translated_paragraphs if p])} 段",
                    info_style,
                )
            )

            # 生成PDF
            doc.build(story)

            # 验证文件是否成功创建
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                return {
                    "status": "success",
                    "output_path": output_path,
                    "format": "mineru_enhanced_pdf",
                }
            else:
                return {"status": "error", "error": "PDF文件生成失败或文件为空"}

        except ImportError as e:
            return {"status": "error", "error": f"缺少必要的依赖库: {str(e)}"}
        except Exception as e:
            return {"status": "error", "error": f"PDF生成过程中出现错误: {str(e)}"}

    def export_translation_result(
        self,
        original_text: str,
        translated_text: str,
        output_path: str,
        format_type: str = "txt",
    ) -> Dict[str, str]:
        """
        导出翻译结果

        Args:
            original_text: 原文
            translated_text: 翻译文本
            output_path: 输出文件路径
            format_type: 输出格式 ('txt', 'docx', 'json', 'pdf')

        Returns:
            导出结果
        """
        try:
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)

            if format_type == "txt":
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write("=== 原文 ===\n")
                    f.write(original_text)
                    f.write("\n\n=== 翻译 ===\n")
                    f.write(translated_text)

            elif format_type == "json":
                import json

                result_data = {
                    "original_text": original_text,
                    "translated_text": translated_text,
                    "timestamp": str(pd.Timestamp.now()),
                    "language": "en_to_zh",
                }
                with open(output_path, "w", encoding="utf-8") as f:
                    json.dump(result_data, f, ensure_ascii=False, indent=2)

            elif format_type == "docx":
                from docx import Document

                doc = Document()
                doc.add_heading("翻译结果", 0)

                doc.add_heading("原文", level=1)
                doc.add_paragraph(original_text)

                doc.add_heading("翻译", level=1)
                doc.add_paragraph(translated_text)

                doc.save(output_path)

            elif format_type == "pdf":
                # 使用新的对比PDF生成方法
                return self.create_comparison_pdf(
                    original_text, translated_text, output_path
                )

            return {
                "status": "success",
                "output_path": output_path,
                "format": format_type,
            }

        except Exception as e:
            return {"status": "error", "error": str(e)}

    def get_pdf_info(self, pdf_path: str) -> Dict[str, any]:
        """
        获取PDF文件信息

        Args:
            pdf_path: PDF文件路径

        Returns:
            PDF信息字典
        """
        try:
            with open(pdf_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)

                info = {
                    "file_path": pdf_path,
                    "file_name": os.path.basename(pdf_path),
                    "total_pages": len(pdf_reader.pages),
                    "file_size": os.path.getsize(pdf_path),
                    "file_size_mb": round(os.path.getsize(pdf_path) / (1024 * 1024), 2),
                }

                # 尝试获取PDF元数据
                if pdf_reader.metadata:
                    info["title"] = pdf_reader.metadata.get("/Title", "未知")
                    info["author"] = pdf_reader.metadata.get("/Author", "未知")
                    info["subject"] = pdf_reader.metadata.get("/Subject", "未知")
                    info["creator"] = pdf_reader.metadata.get("/Creator", "未知")

                return info

        except Exception as e:
            return {"status": "error", "error": str(e), "file_path": pdf_path}

    def is_scanned_pdf(self, pdf_path: str) -> bool:
        """
        判断是否为扫描版PDF

        Args:
            pdf_path: PDF文件路径

        Returns:
            是否为扫描版PDF
        """
        try:
            with pdfplumber.open(pdf_path) as pdf:
                # 检查前几页的文本提取情况
                text_count = 0
                for i in range(min(3, len(pdf.pages))):
                    text = pdf.pages[i].extract_text()
                    if text and len(text.strip()) > 50:  # 如果提取到足够多的文本
                        text_count += 1

                # 如果大部分页面都无法提取到足够文本，可能是扫描版
                return text_count < 2

        except Exception:
            return True

        return False
